#!/usr/bin/env python3
"""Parse quiz questions out of a Word (.docx) document.

Deliberately forgiving, because lecturers' documents are not consistent. Two
layouts are understood:

1. A table whose header row is Question / A / B / C / D / Answer / Explanation.
   Preferred when present — it is far more reliable than prose.
2. Numbered text:

       1. What is the SI unit of force?
       A) Newton
       B) Joule
       C) Watt
       D) Pascal
       Answer: A
       Explanation: Force is measured in newtons.

The awkward part is Word's automatic list numbering: the number is NOT in the
paragraph text python-docx returns. Questions therefore also have to be found by
blank-line grouping and by reading pPr/numPr, not by looking for "1." in text.
That is the single most common way a naive parser fails, so it is tested against
a document that uses real Word auto-numbering.

Nothing here writes a test. It returns JSON for the lecturer to review.
"""

from __future__ import annotations

import io
import re
import unicodedata
from typing import Any

MAX_QUESTIONS = 200
MAX_UPLOAD_BYTES = 5 * 1024 * 1024
REQUIRED_OPTIONS = 4

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# Word's typographic characters, normalised to their plain equivalents.
# Everything else — θ Σ ° µ ± × and the rest — is left exactly as written.
_PUNCTUATION_MAP = {
    "‘": "'", "’": "'", "‚": "'", "‛": "'",
    "“": '"', "”": '"', "„": '"', "‟": '"',
    "–": "-", "—": "-", "‒": "-", "―": "-",
    "…": "...",
    " ": " ", " ": " ", " ": " ", "​": "",
    "﻿": "",
}

_SUPERSCRIPT = {
    "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
    "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹",
    "+": "⁺", "-": "⁻", "(": "⁽", ")": "⁾", "n": "ⁿ",
}
_SUBSCRIPT = {
    "0": "₀", "1": "₁", "2": "₂", "3": "₃", "4": "₄",
    "5": "₅", "6": "₆", "7": "₇", "8": "₈", "9": "₉",
    "+": "₊", "-": "₋", "(": "₍", ")": "₎",
}

OPTION_LETTERS = ["A", "B", "C", "D"]

# "A)" "A." "(A)" "[A]" "a -" ...
OPTION_RE = re.compile(r"^\s*[\(\[]?\s*([A-Ha-h])\s*[\)\].:\-]\s+(.*)$")
# "1." "1)" "Q1." "Question 1:" "Q 1 -"
QUESTION_NUMBER_RE = re.compile(r"^\s*(?:Q(?:uestion)?\s*\.?\s*)?(\d{1,3})\s*[\.\):\-]\s*(.*)$", re.IGNORECASE)
ANSWER_RE = re.compile(r"^\s*(?:correct\s+answer|answer|ans|correct|key)\s*[:\-–]\s*(.+)$", re.IGNORECASE)
EXPLANATION_RE = re.compile(r"^\s*(?:explanation|rationale|feedback|why|reason)\s*[:\-–]\s*(.*)$", re.IGNORECASE)
TIME_RE = re.compile(r"^\s*(?:time|time\s*limit|seconds)\s*[:\-–]\s*(\d{1,3})\s*(?:s|sec|secs|seconds)?\s*$", re.IGNORECASE)

TABLE_QUESTION_HEADERS = {"question", "questions", "question text", "q"}
TABLE_ANSWER_HEADERS = {"answer", "correct", "correct answer", "ans", "key"}
TABLE_EXPLANATION_HEADERS = {"explanation", "rationale", "feedback", "why", "reason"}
TABLE_TIME_HEADERS = {"time", "time limit", "seconds"}


class DocxImportError(ValueError):
    """The file could not be read as a Word document at all."""


# ── Text extraction ──────────────────────────────────────────────────────────

def normalise_text(value: str) -> str:
    if not value:
        return ""
    for bad, good in _PUNCTUATION_MAP.items():
        value = value.replace(bad, good)
    # NFC keeps combining marks together (θ stays θ) without stripping anything.
    value = unicodedata.normalize("NFC", value)
    return re.sub(r"[ \t]+", " ", value).strip()


def _run_text(run) -> str:
    """Run text, mapping formatted super/subscripts to their Unicode forms.

    Word writes "m²" as a plain "2" with vertAlign="superscript"; the character
    itself carries no hint, so the exponent would silently flatten to "m2".
    """
    text = run.text or ""
    if not text:
        return ""
    try:
        superscript = bool(run.font.superscript)
        subscript = bool(run.font.subscript)
    except Exception:
        return text
    if superscript:
        table = _SUPERSCRIPT
    elif subscript:
        table = _SUBSCRIPT
    else:
        return text
    if all(ch in table for ch in text.strip()) and text.strip():
        return "".join(table.get(ch, ch) for ch in text)
    # Mixed content: keep it readable rather than mangling it.
    return f"^({text})" if table is _SUPERSCRIPT else f"_({text})"


def paragraph_text(paragraph) -> str:
    return normalise_text("".join(_run_text(run) for run in paragraph.runs) or paragraph.text)


def paragraph_has_equation(paragraph) -> bool:
    """Word equation objects (OMML) are not part of the text python-docx returns."""
    try:
        return bool(paragraph._element.findall(f".//{{{M_NS}}}oMath")) or \
               bool(paragraph._element.findall(f".//{{{M_NS}}}oMathPara"))
    except Exception:
        return False


def paragraph_is_auto_numbered(paragraph) -> bool:
    """True for Word's automatic list numbering, whose number is not in the text."""
    try:
        p_pr = paragraph._element.find(f"{{{W_NS}}}pPr")
        if p_pr is None:
            return False
        if p_pr.find(f"{{{W_NS}}}numPr") is not None:
            return True
    except Exception:
        return False
    return False


def paragraph_indent_level(paragraph) -> int:
    try:
        p_pr = paragraph._element.find(f"{{{W_NS}}}pPr")
        if p_pr is None:
            return 0
        num_pr = p_pr.find(f"{{{W_NS}}}numPr")
        if num_pr is not None:
            ilvl = num_pr.find(f"{{{W_NS}}}ilvl")
            if ilvl is not None:
                return int(ilvl.get(f"{{{W_NS}}}val") or 0)
    except Exception:
        return 0
    return 0


# ── Question assembly ────────────────────────────────────────────────────────

class _Draft:
    __slots__ = ("q", "options", "answer_raw", "correct_index", "explanation", "time_limit",
                 "has_equation", "closed", "source_index", "explicit")

    def __init__(self, text: str, source_index: int):
        self.q = text
        self.options: list[str] = []
        self.answer_raw: str | None = None
        self.correct_index: int | None = None
        self.explanation = ""
        self.time_limit: int | None = None
        self.has_equation = False
        self.closed = False
        self.source_index = source_index
        # True once something unambiguous marks this as a question: lettered
        # options, an "Answer:" line, or an asterisked option. Without one of
        # those, a run of ordinary prose paragraphs looks exactly like a
        # question with unmarked options, and instruction text or a bulleted
        # list would be imported as a question.
        self.explicit = False


def _strip_correct_marker(text: str) -> tuple[str, bool]:
    """An asterisk before or after an option marks it as the correct one."""
    stripped = text.strip()
    marked = False
    if stripped.startswith("*") and not stripped.startswith("**"):
        stripped = stripped[1:].strip()
        marked = True
    if stripped.endswith("*") and not stripped.endswith("**"):
        stripped = stripped[:-1].strip()
        marked = True
    if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
        stripped = stripped[2:-2].strip()
        marked = True
    return stripped, marked


def _resolve_answer(draft: _Draft) -> tuple[int | None, str | None]:
    """Return (index, warning). Accepts a letter or the option's own text."""
    if draft.correct_index is not None:
        return draft.correct_index, None
    raw = (draft.answer_raw or "").strip()
    if not raw:
        return None, None

    candidate = raw.rstrip(".)]:").strip()
    # "B", "(B)", "Option B", "B) Joule"
    letter_match = re.match(r"^[\(\[]?\s*(?:option\s+)?([A-Ha-h])\s*[\)\].:\-]?\s*(.*)$", candidate)
    if letter_match:
        index = ord(letter_match.group(1).upper()) - ord("A")
        if 0 <= index < len(draft.options):
            return index, None

    # Match by the option's text instead.
    normalised = re.sub(r"\s+", " ", candidate).strip().casefold()
    for index, option in enumerate(draft.options):
        if re.sub(r"\s+", " ", option).strip().casefold() == normalised:
            return index, None
    for index, option in enumerate(draft.options):
        option_norm = re.sub(r"\s+", " ", option).strip().casefold()
        if option_norm and (option_norm in normalised or normalised in option_norm):
            return index, None
    return None, f"Could not match the answer {raw!r} to any option"


def _finalise(draft: _Draft, questions: list[dict[str, Any]], warnings: list[dict[str, Any]], skipped: list[int]) -> None:
    index = draft.source_index

    def warn(message: str) -> None:
        warnings.append({"index": index, "message": message})

    if not draft.explicit:
        # Nothing marked this block as a question. Almost always prose,
        # headings or a bulleted list, so drop it without noise — unless it
        # reads like a question, in which case say why it was not imported.
        if draft.q.strip().endswith("?") and len(draft.options) >= 2:
            warn(
                f"“{draft.q.strip()[:60]}” looks like a question, but the answer options are not "
                "labelled A) B) C) D) and there is no “Answer:” line — skipped."
            )
            skipped.append(index)
        return

    if draft.has_equation:
        warn("This question contains a Word equation that could not be imported — please retype it as text.")

    if not draft.q.strip():
        warn("No question text found — skipped.")
        skipped.append(index)
        return

    options = [opt for opt in (o.strip() for o in draft.options) if opt]
    if len(options) != REQUIRED_OPTIONS:
        warn(
            f"Found {len(options)} answer option{'' if len(options) == 1 else 's'}, "
            f"but a question needs exactly {REQUIRED_OPTIONS} — skipped."
        )
        skipped.append(index)
        return

    correct, answer_warning = _resolve_answer(draft)
    if answer_warning:
        warn(answer_warning + " — defaulted to A.")
        correct = 0
    elif correct is None:
        warn("No correct answer marked — defaulted to A.")
        correct = 0

    question: dict[str, Any] = {
        "q": draft.q.strip()[:600],
        "options": [opt[:240] for opt in options],
        "correct": correct,
        "explanation": draft.explanation.strip()[:2000],
    }
    if draft.time_limit is not None:
        if 5 <= draft.time_limit <= 300:
            question["time_limit"] = draft.time_limit
        else:
            warn(f"Time limit {draft.time_limit}s is outside 5-300 seconds — using the test default.")
    questions.append(question)


# ── Paragraph (numbered-text) parsing ────────────────────────────────────────

def _parse_paragraphs(paragraphs) -> tuple[list[dict[str, Any]], list[dict[str, Any]], int]:
    questions: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    skipped: list[int] = []

    current: _Draft | None = None
    question_counter = 0
    saw_blank = True

    def flush() -> None:
        nonlocal current
        if current is not None:
            _finalise(current, questions, warnings, skipped)
            current = None

    for paragraph in paragraphs:
        text = paragraph_text(paragraph)
        has_equation = paragraph_has_equation(paragraph)
        auto_numbered = paragraph_is_auto_numbered(paragraph)
        indent = paragraph_indent_level(paragraph)

        if not text:
            if has_equation and current is not None:
                current.has_equation = True
            saw_blank = True
            continue

        if len(questions) >= MAX_QUESTIONS:
            break

        # Answer / explanation / time lines belong to the question in progress.
        answer_match = ANSWER_RE.match(text)
        if answer_match and current is not None:
            current.answer_raw = answer_match.group(1).strip()
            current.explicit = True
            current.closed = True
            saw_blank = False
            continue

        explanation_match = EXPLANATION_RE.match(text)
        if explanation_match and current is not None:
            addition = explanation_match.group(1).strip()
            current.explanation = f"{current.explanation} {addition}".strip()
            current.closed = True
            if has_equation:
                current.has_equation = True
            saw_blank = False
            continue

        time_match = TIME_RE.match(text)
        if time_match and current is not None:
            current.time_limit = int(time_match.group(1))
            saw_blank = False
            continue

        # An explicitly marked option, e.g. "B) Joule".
        option_match = OPTION_RE.match(text)
        if option_match and current is not None and not current.closed and len(current.options) < 8:
            body, marked = _strip_correct_marker(option_match.group(2))
            expected_letter = option_match.group(1).upper()
            # Only accept it if the letter follows on from the options so far;
            # otherwise "A. Newton" as a question body would be misread.
            if ord(expected_letter) - ord("A") == len(current.options) or len(current.options) < REQUIRED_OPTIONS:
                if marked:
                    current.correct_index = len(current.options)
                current.explicit = True          # lettered options are unambiguous
                current.options.append(body)
                if has_equation:
                    current.has_equation = True
                saw_blank = False
                continue

        numbered_match = QUESTION_NUMBER_RE.match(text)
        explicit_new_question = bool(numbered_match) and not option_match

        # An unmarked option: Word bullets or auto-numbered lists carry no
        # letter at all, so fall back to position — indented, or directly under
        # a question that still needs options.
        looks_like_unmarked_option = (
            current is not None
            and not current.closed
            and not explicit_new_question
            and len(current.options) < REQUIRED_OPTIONS
            and bool(current.q)
            and (auto_numbered or indent > 0 or not saw_blank)
        )

        if explicit_new_question:
            flush()
            question_counter += 1
            body = numbered_match.group(2).strip()
            current = _Draft(body, question_counter)
            current.has_equation = has_equation
            saw_blank = False
            continue

        if looks_like_unmarked_option:
            body, marked = _strip_correct_marker(text)
            if marked:
                current.correct_index = len(current.options)
                current.explicit = True      # an asterisk marks the correct option
            current.options.append(body)
            if has_equation:
                current.has_equation = True
            saw_blank = False
            continue

        # Continuation of the current question's text (a wrapped line).
        if current is not None and not current.closed and not current.options and not saw_blank:
            current.q = f"{current.q} {text}".strip()
            if has_equation:
                current.has_equation = True
            continue

        # Anything else starts a new question. This is the path Word's automatic
        # numbering takes, since its numbers never reach the paragraph text.
        flush()
        question_counter += 1
        current = _Draft(text, question_counter)
        current.has_equation = has_equation
        saw_blank = False

    flush()
    return questions, warnings, len(skipped)


# ── Table parsing ────────────────────────────────────────────────────────────

def _header_map(table) -> dict[str, int] | None:
    if not table.rows:
        return None
    headers = [normalise_text(cell.text).casefold() for cell in table.rows[0].cells]
    mapping: dict[str, int] = {}
    for index, header in enumerate(headers):
        if header in TABLE_QUESTION_HEADERS and "question" not in mapping:
            mapping["question"] = index
        elif header in TABLE_ANSWER_HEADERS and "answer" not in mapping:
            mapping["answer"] = index
        elif header in TABLE_EXPLANATION_HEADERS and "explanation" not in mapping:
            mapping["explanation"] = index
        elif header in TABLE_TIME_HEADERS and "time" not in mapping:
            mapping["time"] = index
        else:
            letter = header.replace("option", "").strip(" .)(:")
            if letter.upper() in OPTION_LETTERS and f"opt{letter.upper()}" not in mapping:
                mapping[f"opt{letter.upper()}"] = index
    needed = {"question", "answer", "optA", "optB", "optC", "optD"}
    return mapping if needed <= set(mapping) else None


def _parse_tables(tables) -> tuple[list[dict[str, Any]], list[dict[str, Any]], int]:
    questions: list[dict[str, Any]] = []
    warnings: list[dict[str, Any]] = []
    skipped = 0
    counter = 0

    for table in tables:
        mapping = _header_map(table)
        if not mapping:
            continue
        for row in table.rows[1:]:
            if len(questions) >= MAX_QUESTIONS:
                break
            cells = row.cells
            counter += 1

            def cell(key: str, cells=cells, mapping=mapping) -> str:
                index = mapping.get(key)
                if index is None or index >= len(cells):
                    return ""
                return normalise_text(cells[index].text)

            question_text = cell("question")
            if not question_text:
                continue

            draft = _Draft(question_text, counter)
            draft.explicit = True
            for letter in OPTION_LETTERS:
                body, marked = _strip_correct_marker(cell(f"opt{letter}"))
                if marked:
                    draft.correct_index = len(draft.options)
                draft.options.append(body)
            draft.answer_raw = cell("answer")
            draft.explanation = cell("explanation")
            time_text = cell("time")
            if time_text:
                digits = re.search(r"\d{1,3}", time_text)
                if digits:
                    draft.time_limit = int(digits.group(0))
            draft.has_equation = any(
                paragraph_has_equation(p)
                for index in mapping.values()
                if index < len(cells)
                for p in cells[index].paragraphs
            )

            before = len(questions)
            _finalise(draft, questions, warnings, [])
            if len(questions) == before:
                skipped += 1

    return questions, warnings, skipped


# ── Entry point ──────────────────────────────────────────────────────────────

def parse_docx(data: bytes) -> dict[str, Any]:
    """Parse a .docx into review-ready questions. Never writes anything."""
    if len(data) > MAX_UPLOAD_BYTES:
        raise DocxImportError(f"That file is larger than {MAX_UPLOAD_BYTES // (1024 * 1024)} MB.")
    try:
        import docx
    except ImportError as exc:                                  # pragma: no cover
        raise DocxImportError("python-docx is not installed on the server.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise DocxImportError(
            "That file could not be opened as a Word document. Save it as .docx (not .doc or .pdf) and try again."
        ) from exc

    title = ""
    for paragraph in document.paragraphs:
        text = paragraph_text(paragraph)
        if text:
            title = text[:140]
            break

    # A table layout is far more reliable than prose, so it wins on conflict —
    # but a document can legitimately contain both (the template does), so
    # always run both passes and merge.
    table_questions, table_warnings, table_skipped = _parse_tables(document.tables)
    text_questions, text_warnings, text_skipped = _parse_paragraphs(document.paragraphs)

    def key(question: dict[str, Any]) -> str:
        return re.sub(r"\s+", " ", question["q"]).strip().casefold()

    seen = {key(q) for q in table_questions}
    merged_text = [q for q in text_questions if key(q) not in seen]

    questions = table_questions + merged_text
    warnings = table_warnings + (text_warnings if merged_text or not table_questions else [])
    skipped = table_skipped + (text_skipped if merged_text or not table_questions else 0)
    if table_questions and merged_text:
        source = "mixed"
    elif table_questions:
        source = "table"
    else:
        source = "text"

    truncated = False
    if len(questions) > MAX_QUESTIONS:
        questions = questions[:MAX_QUESTIONS]
        truncated = True
        warnings.append({"index": None, "message": f"Only the first {MAX_QUESTIONS} questions were imported."})

    return {
        "questions": questions,
        "warnings": warnings,
        "meta": {
            "title": title,
            "parsed": len(questions),
            "skipped": skipped,
            "layout": source,
            "truncated": truncated,
        },
    }


# ── Downloadable template ────────────────────────────────────────────────────

TEMPLATE_EXAMPLES = [
    {
        "q": "A force of 250 N acts at an angle θ = 30° to the horizontal. What is its horizontal component?",
        "options": ["125 N", "216.5 N", "250 N", "433 N"],
        "answer": "B",
        "explanation": "Fₓ = F cos θ = 250 × cos 30° = 216.5 N.",
        "time": 60,
    },
    {
        "q": "For a body in static equilibrium, which condition must hold?",
        "options": ["ΣF = 0 only", "ΣM = 0 only", "ΣF = 0 and ΣM = 0", "ΣF ± ΣM = 0"],
        "answer": "C",
        "explanation": "Both the resultant force and the resultant moment must vanish.",
        "time": None,
    },
    {
        "q": "What is the SI unit of the second moment of area?",
        "options": ["m²", "m³", "m⁴", "kg·m²"],
        "answer": "m⁴",
        "explanation": "The second moment of area has units of length to the fourth power.",
        "time": None,
    },
]


def build_template_docx() -> bytes:
    """Generate the question template at request time (no binary in the repo)."""
    import docx
    from docx.shared import Pt

    document = docx.Document()
    document.add_heading("Engineering Quiz — question import template", level=0)

    document.add_paragraph(
        "Fill in your questions using either layout below, save the file as .docx, "
        "then upload it with “Import questions from Word” in the test editor. "
        "Nothing is saved until you review the questions and press Save."
    )

    document.add_heading("Rules", level=1)
    for rule in [
        "Every question needs exactly 4 answer options. Questions with 2 or 3 are reported and skipped.",
        "Mark the correct answer with an “Answer:” line, or put an asterisk * on the correct option.",
        "“Answer: B” and “Answer: Newton” both work — the letter or the option's own text.",
        "“Explanation:” is optional and is shown to students after they answer.",
        "“Time: 60” is optional and sets the seconds allowed for that one question (5-300).",
        "Word's automatic numbering is fine. So is typing the numbers yourself, or using none at all.",
        "Symbols such as θ, Σ, °, µ, ±, × and m⁴ import correctly.",
    ]:
        document.add_paragraph(rule, style="List Bullet")

    warning = document.add_paragraph()
    run = warning.add_run(
        "Please note: equations built with Word's equation editor cannot be imported. "
        "Type those as ordinary text instead — for example “F = m a” or “σ = F / A”. "
        "Any question containing one is flagged for you during review."
    )
    run.bold = True

    document.add_heading("Layout 1 — numbered text (recommended)", level=1)
    document.add_paragraph("Copy this pattern. A blank line between questions helps but is not required.")

    for number, example in enumerate(TEMPLATE_EXAMPLES, 1):
        document.add_paragraph(f"{number}. {example['q']}")
        for letter, option in zip(OPTION_LETTERS, example["options"], strict=False):
            document.add_paragraph(f"{letter}) {option}")
        document.add_paragraph(f"Answer: {example['answer']}")
        document.add_paragraph(f"Explanation: {example['explanation']}")
        if example["time"]:
            document.add_paragraph(f"Time: {example['time']}")
        document.add_paragraph("")

    document.add_heading("Layout 2 — table", level=1)
    document.add_paragraph(
        "A table is the most reliable layout. Keep the header row exactly as shown; "
        "the Explanation and Time columns are optional."
    )

    headers = ["Question", "A", "B", "C", "D", "Answer", "Explanation", "Time"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, heading in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = heading
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    example = TEMPLATE_EXAMPLES[0]
    row = table.add_row().cells
    row[0].text = example["q"]
    for index, option in enumerate(example["options"]):
        row[index + 1].text = option
    row[5].text = example["answer"]
    row[6].text = example["explanation"]
    row[7].text = str(example["time"] or "")

    document.add_paragraph("")
    footer = document.add_paragraph()
    footer.add_run(
        "Delete these examples before uploading, or leave them in and remove the questions "
        "you do not want during the review step."
    ).italic = True
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
