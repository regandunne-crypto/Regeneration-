"""Build the .docx fixtures for the import tests at run time.

Generated rather than committed so the repository holds no binaries, and so the
Word auto-numbering fixture is provably real numbering rather than typed text.
"""

from __future__ import annotations

import io

import docx
from docx.oxml.ns import qn


def _save(document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def typed_numbering() -> bytes:
    """Numbers typed by hand — the format the lecturer prefers."""
    document = docx.Document()
    document.add_paragraph("Chapter 3 — Forces")
    document.add_paragraph("")
    document.add_paragraph("1. What is the SI unit of force?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: A")
    document.add_paragraph("Explanation: Force is measured in newtons.")
    document.add_paragraph("")
    document.add_paragraph("Q2. Which quantity is a vector?")
    document.add_paragraph("(a) Mass")
    document.add_paragraph("(b) Speed")
    document.add_paragraph("(c) Velocity")
    document.add_paragraph("(d) Temperature")
    document.add_paragraph("Ans: c")
    document.add_paragraph("")
    document.add_paragraph("Question 3: A moment is measured in which unit?")
    document.add_paragraph("A. N")
    document.add_paragraph("B. N m")
    document.add_paragraph("C. N/m")
    document.add_paragraph("D. m/N")
    document.add_paragraph("Correct: B")
    document.add_paragraph("Rationale: A moment is force times perpendicular distance.")
    return _save(document)


def word_auto_numbering() -> bytes:
    """Word's automatic list numbering.

    The number is stored in pPr/numPr and never appears in the paragraph text,
    so a parser that looks for "1." in the text finds nothing at all. This is
    the fixture that catches the most common failure.
    """
    document = docx.Document()

    def numbered(text: str, num_id: int, level: int = 0):
        paragraph = document.add_paragraph(text, style="List Number")
        p_pr = paragraph._element.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        ilvl = num_pr.get_or_add_ilvl()
        ilvl.set(qn("w:val"), str(level))
        num = num_pr.get_or_add_numId()
        num.set(qn("w:val"), str(num_id))
        return paragraph

    numbered("What is the SI unit of pressure?", num_id=1)
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Pascal")
    document.add_paragraph("C) Joule")
    document.add_paragraph("D) Watt")
    document.add_paragraph("Answer: B")

    numbered("Which law relates force, mass and acceleration?", num_id=1)
    document.add_paragraph("A) Hooke's law")
    document.add_paragraph("B) Newton's second law")
    document.add_paragraph("C) Boyle's law")
    document.add_paragraph("D) Ohm's law")
    document.add_paragraph("Answer: B")
    return _save(document)


def word_auto_numbered_options() -> bytes:
    """Auto-numbered questions AND auto-numbered options: no markers anywhere."""
    document = docx.Document()

    def listed(text: str, level: int):
        paragraph = document.add_paragraph(text, style="List Number")
        p_pr = paragraph._element.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), str(level))
        num_pr.get_or_add_numId().set(qn("w:val"), "2")
        return paragraph

    listed("How many degrees of freedom does a planar rigid body have?", 0)
    listed("One", 1)
    listed("Two", 1)
    listed("Three", 1)
    listed("Six", 1)
    document.add_paragraph("Answer: Three")
    return _save(document)


def no_numbering_at_all() -> bytes:
    """Plain paragraphs separated by blank lines."""
    document = docx.Document()
    document.add_paragraph("What is the SI unit of energy?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: B")
    document.add_paragraph("")
    document.add_paragraph("What is the SI unit of power?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: C")
    return _save(document)


def table_layout() -> bytes:
    document = docx.Document()
    document.add_paragraph("Statics tutorial")
    headers = ["Question", "A", "B", "C", "D", "Answer", "Explanation", "Time"]
    rows = [
        ["What is the SI unit of force?", "Newton", "Joule", "Watt", "Pascal", "A", "Newtons.", "30"],
        ["Sum of moments at equilibrium?", "1", "0", "-1", "Undefined", "B", "ΣM = 0.", "90"],
    ]
    table = document.add_table(rows=1, cols=len(headers))
    for index, heading in enumerate(headers):
        table.rows[0].cells[index].text = heading
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    return _save(document)


def asterisk_marked_answers() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. Which material property describes stiffness?")
    document.add_paragraph("A) Density")
    document.add_paragraph("B) *Young's modulus*")
    document.add_paragraph("C) Hardness")
    document.add_paragraph("D) Toughness")
    document.add_paragraph("")
    document.add_paragraph("2. Which of these is a scalar?")
    document.add_paragraph("A) Force")
    document.add_paragraph("B) Velocity")
    document.add_paragraph("C) Mass*")
    document.add_paragraph("D) Acceleration")
    return _save(document)


def answer_by_text() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. What is the SI unit of force?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: Newton")
    return _save(document)


def three_options_only() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. Name the three states of matter.")
    document.add_paragraph("A) Solid")
    document.add_paragraph("B) Liquid")
    document.add_paragraph("C) Gas")
    document.add_paragraph("Answer: A")
    document.add_paragraph("")
    document.add_paragraph("2. What is the SI unit of force?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: A")
    return _save(document)


def with_equation_object() -> bytes:
    """A real OMML equation, which python-docx text extraction cannot see."""
    from docx.oxml import parse_xml

    document = docx.Document()
    paragraph = document.add_paragraph("1. Evaluate the bending stress using ")
    omml = parse_xml(
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<m:r><m:t>sigma = My/I</m:t></m:r></m:oMath>'
    )
    paragraph._element.append(omml)
    document.add_paragraph("A) 12 MPa")
    document.add_paragraph("B) 24 MPa")
    document.add_paragraph("C) 36 MPa")
    document.add_paragraph("D) 48 MPa")
    document.add_paragraph("Answer: B")
    return _save(document)


def unicode_symbols() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. A force acts at θ = 30° with Σ F = ±250 µN × 2. What is the second moment of area's unit?")
    document.add_paragraph("A) m²")
    document.add_paragraph("B) m³")
    document.add_paragraph("C) m⁴")
    document.add_paragraph("D) kg·m²")
    document.add_paragraph("Answer: C")
    document.add_paragraph("Explanation: “Second moment of area” – note the smart quotes and en dash.")
    return _save(document)


def formatted_superscript() -> bytes:
    """Superscript applied as formatting, not as a Unicode character."""
    document = docx.Document()
    paragraph = document.add_paragraph("1. The unit of the second moment of area is m")
    run = paragraph.add_run("4")
    run.font.superscript = True
    paragraph.add_run(". True or false?")
    document.add_paragraph("A) True")
    document.add_paragraph("B) False")
    document.add_paragraph("C) Only for beams")
    document.add_paragraph("D) Only for shafts")
    document.add_paragraph("Answer: A")
    return _save(document)


def no_answer_marked() -> bytes:
    """Lettered options but no Answer line and no asterisk."""
    document = docx.Document()
    document.add_paragraph("1. What is the SI unit of energy?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    return _save(document)


def out_of_range_time() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. Calculate the reaction at support A.")
    document.add_paragraph("A) 10 kN")
    document.add_paragraph("B) 20 kN")
    document.add_paragraph("C) 30 kN")
    document.add_paragraph("D) 40 kN")
    document.add_paragraph("Answer: C")
    document.add_paragraph("Time: 999")
    return _save(document)


def prose_that_is_not_a_quiz() -> bytes:
    """Headings and bullet lists must not be imported as questions."""
    document = docx.Document()
    document.add_heading("Course notes", level=1)
    document.add_paragraph("These notes cover the topics below and are not a quiz.")
    document.add_paragraph("Equilibrium of rigid bodies", style="List Bullet")
    document.add_paragraph("Free body diagrams", style="List Bullet")
    document.add_paragraph("Distributed loads", style="List Bullet")
    document.add_paragraph("Friction on inclined planes", style="List Bullet")
    return _save(document)


def with_time_lines() -> bytes:
    document = docx.Document()
    document.add_paragraph("1. Quick recall: unit of force?")
    document.add_paragraph("A) Newton")
    document.add_paragraph("B) Joule")
    document.add_paragraph("C) Watt")
    document.add_paragraph("D) Pascal")
    document.add_paragraph("Answer: A")
    document.add_paragraph("Time: 10")
    document.add_paragraph("")
    document.add_paragraph("2. Calculate the reaction at support A for the beam shown.")
    document.add_paragraph("A) 10 kN")
    document.add_paragraph("B) 20 kN")
    document.add_paragraph("C) 30 kN")
    document.add_paragraph("D) 40 kN")
    document.add_paragraph("Answer: C")
    document.add_paragraph("Time: 120")
    return _save(document)
